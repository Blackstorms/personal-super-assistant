#!/usr/bin/env python3
"""将 docs/项目说明文档.md 转为 Word（.docx）。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "项目说明文档.md"
OUT = ROOT / "docs" / "项目说明文档.docx"


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def add_runs_with_code(paragraph, content: str, *, base_size=Pt(12)):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", content)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, east_asia="宋体", ascii_font="Consolas", size=Pt(10.5))
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    body = s.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")
    return body == ""


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "黑体"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        hs.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)

    i = 0
    in_code = False
    code_lines: list[str] = []
    title_done = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.3)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, ascii_font="Consolas", east_asia="宋体", size=Pt(9))
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            if level == 1 and not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title)
                set_run_font(run, east_asia="黑体", ascii_font="Arial", size=Pt(22), bold=True)
                title_done = True
            else:
                doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_sep(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    val = row[c_idx] if c_idx < len(row) else ""
                    add_runs_with_code(p, val, base_size=Pt(10.5))
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_code(p, m.group(2))
            i += 1
            continue

        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_code(p, line.lstrip()[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_code(p, line.strip())
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
